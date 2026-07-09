from __future__ import annotations

from pathlib import Path

import pytest

from alpi import extract as ext
from alpi.tools import workspace as ws


def test_vec_blob_roundtrips_float_count() -> None:
    blob = ws._vec_blob([1.0, 2.0, 3.0])
    assert len(blob) == 12


def test_chunk_lines_is_zero_based_stride_source() -> None:
    body = "\n".join(f"line {i}" for i in range(120))
    rows = ws._chunk_lines(body)
    assert len(rows) >= 4
    assert rows[0][0] == 1
    assert rows[0][1] == 30
    assert rows[1][0] == 26


def test_reader_dispatcher_routes_by_suffix() -> None:
    assert ws._reader_for(".md") is ws._read_text
    assert ws._reader_for(".py") is ws._read_text
    assert ws._reader_for(".html") is ws._read_html
    assert ws._reader_for(".htm") is ws._read_html
    assert ws._reader_for(".docx") is ws._read_docx
    assert ws._reader_for(".epub") is ws._read_epub
    pdf_reader = ws._reader_for(".pdf")
    assert callable(pdf_reader) and pdf_reader is not ws._read_pdf
    img_reader = ws._reader_for(".jpg")
    assert callable(img_reader) and img_reader is not ws._read_image


def test_read_text_replaces_invalid_bytes(tmp_path: Path) -> None:
    path = tmp_path / "bad.txt"
    path.write_bytes(b"hello\xffworld")
    assert "hello" in ws._read_text(path)


def test_read_docx(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("This document describes quarterly revenue.")
    doc.add_paragraph("")
    doc.add_paragraph("Q3 closed ahead of forecast.")
    path = tmp_path / "report.docx"
    doc.save(path)

    text = ws._read_docx(path)
    assert "quarterly revenue" in text
    assert "ahead of forecast" in text


def test_pdf_with_text_layer_skips_ocr(monkeypatch, tmp_path: Path) -> None:
    pdf = tmp_path / "doc.pdf"
    pdf.write_bytes(b"placeholder")

    monkeypatch.setattr(
        ext,
        "ocr_pdf",
        lambda p, **kw: pytest.fail("OCR fallback should not run for text-layer PDFs"),
    )

    class FakePage:
        def extract_text(self):
            return "lots of real text " * 10

    class FakeReader:
        pages = [FakePage(), FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", lambda _: FakeReader())
    out = ws._read_pdf(pdf, ocr=False)
    assert "real text" in out


def test_pdf_without_text_layer_requires_ocr_flag(monkeypatch, tmp_path: Path) -> None:
    pdf = tmp_path / "scanned.pdf"
    pdf.write_bytes(b"placeholder")

    class EmptyPage:
        def extract_text(self):
            return ""

    class EmptyReader:
        pages = [EmptyPage()]

    monkeypatch.setattr("pypdf.PdfReader", lambda _: EmptyReader())

    with pytest.raises(ws.OcrRequired) as err:
        ws._read_pdf(pdf, ocr=False)
    assert 'knowledge(action="ingest", ocr=true)' in str(err.value)


def test_pdf_without_text_layer_falls_back_to_ocr_when_flag_on(
    monkeypatch, tmp_path: Path,
) -> None:
    pdf = tmp_path / "scanned.pdf"
    pdf.write_bytes(b"placeholder")
    called = {"ocr": False}

    def fake_ocr(p, **kw):
        called["ocr"] = True
        return ("OCR-extracted text from the scan", False)

    monkeypatch.setattr(ext, "ocr_pdf", fake_ocr)

    class EmptyPage:
        def extract_text(self):
            return ""

    class EmptyReader:
        pages = [EmptyPage()]

    monkeypatch.setattr("pypdf.PdfReader", lambda _: EmptyReader())
    out = ws._read_pdf(pdf, ocr=True)
    assert "OCR-extracted" in out
    assert called["ocr"] is True


def test_image_requires_ocr_flag(tmp_path: Path) -> None:
    image = tmp_path / "receipt.jpg"
    image.write_bytes(b"fake")
    with pytest.raises(ws.OcrRequired) as err:
        ws._read_image(image, ocr=False)
    assert 'knowledge(action="ingest", ocr=true)' in str(err.value)


def test_image_path_uses_ocr_when_flag_on(tmp_path: Path, monkeypatch) -> None:
    from PIL import Image

    image = tmp_path / "receipt.jpg"
    Image.new("RGB", (2, 2), color="white").save(image)

    monkeypatch.setattr(ext, "ocr_image", lambda p: "Receipt total 4.50 EUR")
    assert "Receipt total" in ws._read_image(image, ocr=True)
